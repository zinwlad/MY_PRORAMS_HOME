import fitz  # PyMuPDF
import os
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from docx import Document

def points_to_mm(points):
    """Преобразует размеры из точек в миллиметры."""
    return points * 25.4 / 72

def analyze_pdf(pdf_path):
    """Анализирует PDF-файл и собирает статистику."""
    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        return f"Не удалось открыть PDF-файл: {e}"

    total_pages = doc.page_count
    rgb_pages = set()
    cmyk_pages = set()
    grayscale_pages = set()
    font_set = set()
    font_sizes = []
    page_sizes_mm = []
    text_blocks_count = 0
    words_count = 0
    chars_count = 0
    image_stats = {'RGB': 0, 'CMYK': 0, 'Grayscale': 0}
    annotation_count = 0
    link_count = 0
    total_images = 0
    page_image_count = {}
    image_formats = set()
    metadata = {
        "author": "",
        "creator": "",
        "created_date": "",
        "modified_date": "",
        "title": ""
    }

    # Чтение метаданных
    try:
        metadata['author'] = doc.metadata.get('author', 'Неизвестно')
        metadata['creator'] = doc.metadata.get('creator', 'Неизвестно')
        metadata['created_date'] = doc.metadata.get('creationDate', 'Неизвестно')
        metadata['modified_date'] = doc.metadata.get('modDate', 'Неизвестно')
        metadata['title'] = doc.metadata.get('title', 'Неизвестно')
    except Exception as e:
        pass

    for page_num in range(total_pages):
        try:
            page = doc.load_page(page_num)
        except Exception as e:
            continue

        page_size = page.rect
        width_mm = points_to_mm(page_size.width)
        height_mm = points_to_mm(page_size.height)
        page_sizes_mm.append((width_mm, height_mm))

        # Подсчет аннотаций и ссылок
        annotations = list(page.annots())  # Преобразование генератора в список
        annotation_count += len(annotations)
        for link in page.get_links():
            if link.get('uri'):
                link_count += 1

        try:
            page_dict = page.get_text("dict")
            if 'fonts' in page_dict:
                fonts = page_dict['fonts']
                for font in fonts:
                    font_name = font['fontname']
                    font_size = font.get('size', 0)
                    font_set.add(font_name)
                    font_sizes.append(font_size)
            if 'blocks' in page_dict:
                text_blocks_count += len(page_dict['blocks'])
                for block in page_dict['blocks']:
                    if block['type'] == 0:  # Text block
                        for line in block['lines']:
                            words_count += len(line['spans'])
                            chars_count += sum(len(span['text']) for span in line['spans'])
        except Exception as e:
            pass

        # Подсчет изображений
        try:
            images = page.get_images(full=True)
            page_image_count[page_num + 1] = len(images)  # Количество изображений на странице
            total_images += len(images)

            if images:
                for img in images:
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    if pix.colorspace.n == 3:
                        rgb_pages.add(page_num + 1)
                        image_stats['RGB'] += 1
                    elif pix.colorspace.n == 4:
                        cmyk_pages.add(page_num + 1)
                        image_stats['CMYK'] += 1
                    elif pix.colorspace.n == 1:  # Assuming grayscale
                        grayscale_pages.add(page_num + 1)
                        image_stats['Grayscale'] += 1
                    # Определение формата изображения
                    image_format = pix.alpha * 8  # Условно определяем формат (это не совсем точный метод)
                    if image_format == 24:
                        image_formats.add('RGB')
                    elif image_format == 32:
                        image_formats.add('CMYK')
                    else:
                        image_formats.add('Unknown')
        except Exception as e:
            pass

    doc.close()

    # Определяем общий формат страниц
    if all(size == page_sizes_mm[0] for size in page_sizes_mm):
        page_format = f"{page_sizes_mm[0][0]:.1f} x {page_sizes_mm[0][1]:.1f} мм"
    else:
        page_format = "Разные размеры страниц"

    # Формирование результата
    result = []
    result.append(f"Количество страниц: {total_pages}")
    result.append(f"Формат страниц: {page_format}")

    result.append(f"\nОбщее количество изображений: {total_images}")
    result.append("\nRGB изображения находятся на страницах:")
    if rgb_pages:
        result.append(", ".join(map(str, sorted(rgb_pages))))
    else:
        result.append("Нет изображений в RGB цвете.")

    result.append("\nCMYK изображения находятся на страницах:")
    if cmyk_pages:
        result.append(", ".join(map(str, sorted(cmyk_pages))))
    else:
        result.append("Нет изображений в CMYK цвете.")

    result.append("\nGrayscale изображения находятся на страницах:")
    if grayscale_pages:
        result.append(", ".join(map(str, sorted(grayscale_pages))))
    else:
        result.append("Нет изображений в Grayscale цвете.")

    result.append("\nИспользованные шрифты:")
    if font_set:
        for font in sorted(font_set):
            result.append(f"  - {font}")
    else:
        result.append("Шрифты отсутствуют или недоступны.")

    result.append(f"\nКоличество текстовых блоков: {text_blocks_count}")
    result.append(f"Количество слов: {words_count}")
    result.append(f"Количество символов: {chars_count}")

    result.append("\nСтатистика по изображениям:")
    for color_space, count in image_stats.items():
        result.append(f"  - {color_space}: {count}")

    if font_sizes:
        result.append(f"\nМинимальный размер шрифта: {min(font_sizes)}")
        result.append(f"Максимальный размер шрифта: {max(font_sizes)}")

    result.append(f"\nКоличество аннотаций: {annotation_count}")
    result.append(f"Количество ссылок: {link_count}")

    result.append("\nМетаданные:")
    result.append(f"  Автор: {metadata['author']}")
    result.append(f"  Создатель: {metadata['creator']}")
    result.append(f"  Дата создания: {metadata['created_date']}")
    result.append(f"  Дата модификации: {metadata['modified_date']}")
    result.append(f"  Титул: {metadata['title']}")

    result.append("\nФорматы изображений:")
    if image_formats:
        result.append(", ".join(sorted(image_formats)))
    else:
        result.append("Форматы изображений не определены.")

    return "\n".join(result)

def analyze_docx(docx_path):
    """Анализирует DOCX-файл и собирает статистику."""
    try:
        doc = Document(docx_path)
    except Exception as e:
        return f"Не удалось открыть DOCX-файл: {e}"

    total_paragraphs = len(doc.paragraphs)
    total_runs = sum(len(p.runs) for p in doc.paragraphs)
    total_words = sum(len(p.text.split()) for p in doc.paragraphs)
    total_chars = sum(len(p.text) for p in doc.paragraphs)

    result = []
    result.append(f"Количество абзацев: {total_paragraphs}")
    result.append(f"Количество строк: {total_runs}")
    result.append(f"Количество слов: {total_words}")
    result.append(f"Количество символов: {total_chars}")

    return "\n".join(result)

def analyze_txt(txt_path):
    """Анализирует TXT-файл и собирает статистику."""
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            content = file.read()
    except Exception as e:
        return f"Не удалось открыть TXT-файл: {e}"

    total_lines = content.count('\n')
    total_words = len(content.split())
    total_chars = len(content)

    result = []
    result.append(f"Количество строк: {total_lines}")
    result.append(f"Количество слов: {total_words}")
    result.append(f"Количество символов: {total_chars}")

    return "\n".join(result)

def browse_file():
    """Открывает диалоговое окно для выбора файла и запускает анализ."""
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf"), ("DOCX files", "*.docx"), ("Text files", "*.txt")])
    if file_path:
        extension = os.path.splitext(file_path)[1].lower()
        if extension == '.pdf':
            result_text = analyze_pdf(file_path)
        elif extension == '.docx':
            result_text = analyze_docx(file_path)
        elif extension == '.txt':
            result_text = analyze_txt(file_path)
        else:
            result_text = "Неподдерживаемый формат файла."
        result_text_widget.config(state=tk.NORMAL)
        result_text_widget.delete(1.0, tk.END)
        result_text_widget.insert(tk.END, result_text)
        result_text_widget.config(state=tk.DISABLED)

def reset():
    """Сбрасывает результаты и очищает текстовое поле."""
    result_text_widget.config(state=tk.NORMAL)
    result_text_widget.delete(1.0, tk.END)
    result_text_widget.config(state=tk.DISABLED)

# Настройка интерфейса
root = tk.Tk()
root.title("Анализатор файлов")

frame = tk.Frame(root)
frame.pack(padx=10, pady=10)

browse_button = tk.Button(frame, text="Открыть файл", command=browse_file)
browse_button.pack(side=tk.LEFT)

reset_button = tk.Button(frame, text="Сброс", command=reset)
reset_button.pack(side=tk.LEFT)

result_text_widget = tk.Text(root, wrap=tk.WORD, height=20, width=80)
result_text_widget.pack(padx=10, pady=10)
result_text_widget.config(state=tk.DISABLED)

# Запуск основного цикла обработки событий
root.mainloop()
