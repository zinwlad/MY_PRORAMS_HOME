from docx import Document

def analyze_docx(docx_path):
    """Анализирует DOCX-файл и собирает расширенную статистику, включая изображения, графики и таблицы."""
    try:
        doc = Document(docx_path)
    except Exception as e:
        return f"Не удалось открыть DOCX-файл: {e}"

    total_paragraphs = len(doc.paragraphs)
    total_runs = 0
    total_words = 0
    total_chars = 0
    bold_count = 0
    italic_count = 0
    underline_count = 0
    heading_count = 0
    table_count = len(doc.tables)
    image_count = 0
    shape_count = 0
    style_stats = {}

    for paragraph in doc.paragraphs:
        total_runs += len(paragraph.runs)
        total_words += len(paragraph.text.split())
        total_chars += len(paragraph.text)

        # Подсчет заголовков
        if paragraph.style.name.startswith('Heading'):
            heading_count += 1

        # Сбор статистики по стилям абзацев
        style_name = paragraph.style.name
        if style_name in style_stats:
            style_stats[style_name] += 1
        else:
            style_stats[style_name] = 1

        # Подсчет форматирования текста
        for run in paragraph.runs:
            if run.bold:
                bold_count += 1
            if run.italic:
                italic_count += 1
            if run.underline:
                underline_count += 1

    # Подсчет изображений и фигур
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1

    # Подсчет фигур (включая графики, если они существуют как объекты OLE)
    for shape in doc.inline_shapes:
        shape_count += 1

    # Формирование результата
    result = []
    result.append(f"Общее количество абзацев: {total_paragraphs}")
    result.append(f"Количество строк (runs): {total_runs}")
    result.append(f"Количество слов: {total_words}")
    result.append(f"Количество символов: {total_chars}")
    result.append(f"Количество заголовков: {heading_count}")
    result.append(f"Количество таблиц: {table_count}")
    result.append(f"Количество изображений: {image_count}")
    result.append(f"Количество встроенных фигур (включая графики): {shape_count}")

    result.append("\nСтатистика стилей абзацев:")
    for style, count in style_stats.items():
        result.append(f"  - {style}: {count}")

    result.append("\nСтатистика форматирования текста:")
    result.append(f"  - Жирный текст: {bold_count}")
    result.append(f"  - Курсивный текст: {italic_count}")
    result.append(f"  - Подчеркнутый текст: {underline_count}")

    return "\n".join(result)
